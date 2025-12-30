from __future__ import annotations

import os
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.shared import Pt, Inches

from .barcode_utils import create_barcode_temp

SLOTS_PER_PAGE = 6  # 2 rows x 3 cols

def _clear_cell(cell):
    cell.text = ""
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

def _add_centered_line(cell, text: str, size: int, bold: bool = False):
    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold

def _add_barcode(cell, code: str):
    """
    Adds barcode image to the cell (centered).
    """
    tmp = create_barcode_temp(code)
    try:
        p = cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run()
        r.add_picture(tmp, width=Inches(2.2))
    finally:
        try:
            os.remove(tmp)
        except OSError:
            pass

def fill_label_cell(cell, lot_number: str, wo: dict, *, sheet_number: int | None = None,
                    qty_override: int | None = None, hide_qty: bool = False):
    """
    Standard label cell layout.
    - cover: qty_override = total_qty
    - sheet labels: hide_qty = True, and include SHEET line
    """
    _clear_cell(cell)

    # Optional sheet line
    if sheet_number is not None:
        _add_centered_line(cell, f"SHEET {sheet_number}", size=16, bold=True)

    # Part
    part = str(wo.get("part", "")).strip()
    _add_centered_line(cell, f"PART {part}", size=14, bold=True)

    # Tag + description
    tag_desc = str(wo.get("tag_desc", "")).strip()
    if tag_desc:
        _add_centered_line(cell, tag_desc, size=11, bold=False)

    # Barcode
    code = str(wo.get("code", "")).strip()
    if code:
        _add_barcode(cell, code)

    # WO number
    wo_num = str(wo.get("work_order", "")).strip()
    if wo_num:
        _add_centered_line(cell, f"WO {wo_num}", size=12, bold=True)

    # LOT number
    _add_centered_line(cell, f"LOT {lot_number}", size=12, bold=True)

    # Quantity line
    if not hide_qty:
        qty_val = qty_override if qty_override is not None else 1
        _add_centered_line(cell, f"QTY {qty_val}", size=12, bold=True)
