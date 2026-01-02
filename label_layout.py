import os
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.shared import Pt, Inches

from barcode_utils import create_barcode_temp


# -------------------------------------------------
# CELL UTILITIES
# -------------------------------------------------

def clear_cell(cell):
    """
    Clear all content from a table cell and vertically center it.
    """
    cell.text = ""
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def add_centered_text(cell, text, size=10, bold=False):
    """
    Add a centered paragraph with Calibri font.

    Parameters:
        cell: python-docx table cell
        text (str): Text to insert
        size (int): Font size in points
        bold (bool): Whether text should be bold
    """
    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Remove extra spacing to keep labels compact
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0

    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.bold = bold


# -------------------------------------------------
# COVER PAGE LABELS
# -------------------------------------------------
def add_cover_label(cell, lot_number):
    """
    Cover slot format:
      Cover Page
      LOT # <lot_number>
    """
    clear_cell(cell)
    add_centered_text(cell, "Cover Page", size=24, bold=True)
    add_centered_text(cell, f"LOT # {lot_number}", size=24, bold=True)



def add_sheet_label(cell, text):
    """
    Sheet slot format.
    The text is already pre-formatted by main.py.

    Example received text:
      "1 - LOT # 08500"

    Output:
      Sheet # 1
      LOT # 08500
    """
    clear_cell(cell)

    # Split sheet number and lot
    if "-" in text:
        sheet_part, lot_part = text.split("-", 1)
        sheet_part = sheet_part.strip()
        lot_part = lot_part.strip()
    else:
        sheet_part = text.strip()
        lot_part = ""

    add_centered_text(cell, f"Sheet # {sheet_part}", size=20, bold=True)

    if lot_part:
        add_centered_text(cell, lot_part, size=20, bold=True)




# -------------------------------------------------
# WORK ORDER LABEL
# -------------------------------------------------

def add_workorder_label(cell, wo, lot_number):
    """
    Render a Work Order label.

    Behavior depends on flags inside the `wo` dictionary:

    - Default:
        * Prints PART, TAG+DESC, barcode, WO, LOT, QTY 1
    - Cover Page:
        * If wo["qty_override"] exists → prints QTY total
    - Sheets:
        * If wo["hide_qty"] == True → QTY line is omitted

    Parameters:
        cell: python-docx table cell
        wo (dict): Work order data
        lot_number (str): Current lot number
    """
    clear_cell(cell)

    # PART NUMBER (largest text)
    add_centered_text(cell, wo["part"], size=26, bold=True)

    # TAG + DESCRIPTION (single combined line)
    add_centered_text(cell, wo["tag_desc"], size=20)

    # -------------------------------------------------
    # BARCODE
    # -------------------------------------------------
    barcode_path = create_barcode_temp(wo["code"])

    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)

    run = p.add_run()
    run.add_picture(barcode_path, width=Inches(1.35))

    # Remove temporary barcode image
    os.remove(barcode_path)

    # -------------------------------------------------
    # TEXT FOOTER
    # -------------------------------------------------

    # Work Order number
    add_centered_text(cell, f'WO {wo["work_order"]}', size=16)

    # Lot number
    add_centered_text(cell, f'LOT {lot_number}', size=16)

    # Quantity handling:
    # - Default → QTY 1
    # - Cover → QTY total (via qty_override)
    # - Sheets → no QTY line (via hide_qty)
    if not wo.get("hide_qty", False):
        add_centered_text(
            cell,
            f'QTY {wo.get("qty_override", 1)}',
            size=16
        )
