import os
import time
from docx import Document

from label_layout import (
    add_cover_label,
    add_sheet_label,
    add_workorder_label,
)

# Footer fields for "Page X of Y"
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

os.makedirs("output", exist_ok=True)

# Fixed 6 slots per page (3 rows x 2 cols)
SLOTS = [(0, 0), (0, 1), (1, 0), (1, 1), (2, 0), (2, 1)]


# -----------------------------
# Input helpers (safe input)
# -----------------------------
def input_int(prompt: str, min_val: int = 0, max_val: int | None = None) -> int:
    """Read an integer from the user with validation."""
    while True:
        try:
            v = int(input(prompt))
            if v < min_val:
                raise ValueError
            if max_val is not None and v > max_val:
                raise ValueError
            return v
        except ValueError:
            print("❌ Please enter a valid number.")


def input_text(prompt: str) -> str:
    """Read a non-empty string from the user."""
    while True:
        t = input(prompt).strip()
        if t:
            return t
        print("❌ This field is required.")


def input_yes_no(prompt: str) -> bool:
    """Return True for Y, False for N (keeps asking until valid)."""
    while True:
        ans = input(prompt).strip().upper()
        if ans in ("Y", "N"):
            return ans == "Y"
        print("❌ Please answer Y or N.")


# -----------------------------
# Footer: "Page X of Y"
# -----------------------------
def add_page_x_of_y_footer(document: Document) -> None:
    """Add centered 'Page X of Y' to the footer using Word fields."""
    section = document.sections[0]
    footer = section.footer
    footer.is_linked_to_previous = False

    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = 1  # Center

    p.add_run("Page ")

    fld_page = OxmlElement("w:fldSimple")
    fld_page.set(qn("w:instr"), "PAGE")
    p.runs[-1]._r.append(fld_page)

    p.add_run(" of ")

    fld_pages = OxmlElement("w:fldSimple")
    fld_pages.set(qn("w:instr"), "NUMPAGES")
    p.runs[-1]._r.append(fld_pages)


# -----------------------------
# Work order edit menu
# -----------------------------
def print_workorders(work_orders: list[dict]) -> None:
    """Print current WO data for review."""
    print("\n=== WORK ORDERS ===")
    for i, wo in enumerate(work_orders, start=1):
        print(
            f"{i}) WO {wo['work_order']} | PART: {wo['part']} | TAG+DESC: {wo['tag_desc']} "
            f"| CODE: {wo['code']} | TOTAL QTY: {wo['total_qty']}"
        )


def edit_workorders(work_orders: list[dict]) -> None:
    """
    Allow editing WO fields BEFORE nesting.
    After editing total_qty, remaining is reset to match it.
    """
    while True:
        print_workorders(work_orders)
        print("\nOptions:")
        print("  [E] Edit a Work Order")
        print("  [C] Continue to sheet nesting")
        choice = input("Select: ").strip().upper()

        if choice == "C":
            return

        if choice != "E":
            print("❌ Invalid option.")
            continue

        idx = input_int(f"Which WO do you want to edit (1-{len(work_orders)}): ", 1, len(work_orders)) - 1
        wo = work_orders[idx]

        print("\nFields:")
        print("  1) Part #")
        print("  2) TAG + DESCRIPTION")
        print("  3) Barcode Code")
        print("  4) WO #")
        print("  5) TOTAL QTY")
        field = input_int("Field: ", 1, 5)

        if field == 1:
            wo["part"] = input_text("New Part #: ")
        elif field == 2:
            wo["tag_desc"] = input_text("New TAG + DESCRIPTION: ")
        elif field == 3:
            wo["code"] = input_text("New Barcode Code: ")
        elif field == 4:
            wo["work_order"] = input_text("New WO #: ")
        elif field == 5:
            wo["total_qty"] = input_int("New TOTAL QTY: ", 1)

        # Keep remaining consistent after edits
        wo["remaining"] = wo["total_qty"]
        print("✅ Work Order updated.")


# -----------------------------
# Sheet planning (manual nesting)
# -----------------------------
def reset_remaining(work_orders: list[dict]) -> None:
    """Reset remaining quantities back to totals (used for re-nesting)."""
    for wo in work_orders:
        wo["remaining"] = wo["total_qty"]


def plan_sheets(work_orders: list[dict]) -> list[dict]:
    """
    Manual nesting by sheet.
    You enter how many parts of each WO go into each sheet.
    A sheet ends after you've entered quantities for ALL WOs.
    """
    sheets: list[dict] = []
    sheet_number = 1

    while any(wo["remaining"] > 0 for wo in work_orders):
        print(f"\n=== DEFINE SHEET {sheet_number} ===")
        allocations: list[tuple[int, int]] = []

        # Ask per WO (includes PART name for clarity)
        for i, wo in enumerate(work_orders):
            if wo["remaining"] <= 0:
                allocations.append((i, 0))
                continue

            qty = input_int(
                f"WO {wo['work_order']} | PART {wo['part']} (remaining {wo['remaining']}): ",
                0,
                wo["remaining"]
            )
            allocations.append((i, qty))

        # Apply decrement after the sheet is defined (sheet ends here)
        for i, qty in allocations:
            work_orders[i]["remaining"] -= qty

        sheets.append({"sheet_number": sheet_number, "allocations": allocations})
        sheet_number += 1

    return sheets


def show_summary(work_orders: list[dict], sheets: list[dict]) -> None:
    """Print a summary of all sheets and totals per WO."""
    print("\n================= NEST SUMMARY =================")
    totals = {i: 0 for i in range(len(work_orders))}

    for sh in sheets:
        print(f"\nSHEET {sh['sheet_number']}:")
        for i, qty in sh["allocations"]:
            if qty > 0:
                wo = work_orders[i]
                totals[i] += qty
                print(f"  - WO {wo['work_order']} | PART {wo['part']} -> {qty} pcs")

    print("\nTOTALS BY WORK ORDER:")
    for i, wo in enumerate(work_orders):
        print(f"  WO {wo['work_order']} | PART {wo['part']} -> {totals[i]} pcs (expected {wo['total_qty']})")


# -----------------------------
# DOC generation (uses your existing label functions)
# -----------------------------
def generate_doc(lot_number: str, work_orders: list[dict], sheets: list[dict]) -> str:
    """
    Build the DOCX using your current formatting functions.
    - Cover page: LOT + up to 4 WOs, QTY = total (override)
    - Sheets: each SHEET label uses the next available slot (no forced page breaks between sheets)
    - Sheets labels: no QTY line (hide flag)
    """
    doc = Document()

    def new_page():
        table = doc.add_table(rows=3, cols=2)
        return table, 0

    # ---------------- COVER PAGE ----------------
    table, slot_idx = new_page()

    # Slot 0: LOT
    add_cover_label(table.cell(0, 0), lot_number)

    # Slots 1-4: WOs (slot 5 stays blank)
    slot_idx = 1
    for wo in work_orders:
        if slot_idx >= 5:
            break

        r, c = SLOTS[slot_idx]

        # Cover: show total QTY using qty_override (do not remove any existing code in label_layout)
        wo["qty_override"] = wo["total_qty"]
        add_workorder_label(table.cell(r, c), wo, lot_number)
        del wo["qty_override"]

        slot_idx += 1

    # After cover, start next content on a new page
    doc.add_page_break()

    # ---------------- SHEETS (continuous flow) ----------------
    table, slot_idx = new_page()

    for sh in sheets:
        # Place the SHEET label in the next available slot
        if slot_idx >= 6:
            doc.add_page_break()
            table, slot_idx = new_page()

        r, c = SLOTS[slot_idx]
        add_sheet_label(table.cell(r, c), sh["sheet_number"])
        slot_idx += 1

        # Print labels per piece
        for i, qty in sh["allocations"]:
            if qty <= 0:
                continue
            wo = work_orders[i]

            for _ in range(qty):
                if slot_idx >= 6:
                    doc.add_page_break()
                    table, slot_idx = new_page()

                r, c = SLOTS[slot_idx]

                # Sheets: hide QTY line (your add_workorder_label checks this flag)
                wo["hide_qty"] = True
                add_workorder_label(table.cell(r, c), wo, lot_number)
                del wo["hide_qty"]

                slot_idx += 1

    # Footer numbering
    add_page_x_of_y_footer(doc)

    filename = f"output/WorkOrder_Labels_{int(time.time())}.docx"
    doc.save(filename)
    return filename


# -----------------------------
# Main program flow
# -----------------------------
def main():
    # 1) Basic LOT + WO entry
    lot_number = input_text("LOT #: ")
    wo_count = input_int("How many Work Orders (max 4): ", 1, 4)

    work_orders: list[dict] = []
    for i in range(wo_count):
        print(f"\n--- WORK ORDER {i + 1} ---")
        wo = {
            "part": input_text("Part #: "),
            "tag_desc": input_text("TAG + DESCRIPTION (single text): "),
            "code": input_text("Barcode code: "),
            "work_order": input_text("WO #: "),
            "total_qty": input_int("TOTAL QTY (entire LOT): ", 1),
        }
        wo["remaining"] = wo["total_qty"]
        work_orders.append(wo)

    # 2) Optional editing BEFORE nesting
    edit_workorders(work_orders)

    # 3) Nesting loop: plan sheets -> summary -> re-nest?
    while True:
        reset_remaining(work_orders)
        sheets = plan_sheets(work_orders)

        show_summary(work_orders, sheets)

        # Ask if user wants to redo nesting (re-enter sheet allocations)
        redo = input_yes_no("\nDo you want to re-nest and re-enter sheet allocations? (Y/N): ")
        if not redo:
            break

    # 4) Final decision to generate/print
    do_print = input_yes_no("\nGenerate/print the DOCX now? (Y/N): ")
    if not do_print:
        print("✅ Cancelled. No document generated.")
        return

    # 5) Generate final document
    filename = generate_doc(lot_number, work_orders, sheets)
    print(f"\n✅ Document generated:\n{filename}")
    print("Tip: In Word, if needed press Ctrl+A then F9 to update 'Page X of Y'.")


if __name__ == "__main__":
    main()
